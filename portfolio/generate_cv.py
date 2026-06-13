"""Generate a professional, ATS-friendly CV in PDF and Word.

English is the base language; a Dutch version is generated too. Privacy-safe
(no home address, phone or birth date), single column, real selectable text and
standard headings. Run:

    python generate_cv.py    # assets/cv.pdf + cv.docx (EN) and cv_nl.pdf + cv_nl.docx (NL)

Edit CONTENT below and re-run.
"""

from __future__ import annotations

from pathlib import Path

from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate

INK = HexColor("#1f2a44")
ACCENT = HexColor("#2a9d8f")
GREY = HexColor("#5b6675")
RULE = HexColor("#d7dee6")

NAME = "Gerrit Düthler"
CONTACT = "gerrit@duthler.info  ·  Eindhoven, Netherlands  ·  linkedin.com/in/gerrit-d-90a50234"

CONTENT = {
    "en": {
        "role": "Cost Engineer · Estimator · Manufacturing Engineer",
        "headline": "Reliable cost prices · faster quotes · margins you can defend",
        "headings": {"summary": "Profile", "competencies": "Core competencies",
                     "experience": "Professional experience", "education": "Education",
                     "certs": "Certifications", "languages": "Languages"},
        "summary": (
            "Cost engineer with a career from the shop floor to the calculation office — DAF "
            "Trucks, VDL ETG, Andritz, Wilting and Wärtsilä — so every cost price I produce is "
            "grounded in what actually happens at the machine. I build defensible cost prices for "
            "development, customer and non-standard work, keep quoting fast and transparent, and "
            "align engineering, purchasing, sales and business control on the same figure. Where "
            "the tooling falls short, I build it myself — automation and dashboards in Python, "
            "and Excel/VBA. Lean Six Sigma Green Belt."
        ),
        "competencies": (
            "Cost estimating & calculation · Should-cost · RFQ & quotation · Post-calculation · "
            "Make-or-buy · Bill of materials (BOM) · Margin & cost control · Value engineering · "
            "Lean Six Sigma (Green Belt) · Continuous improvement · Work preparation & routing · "
            "Manufacturing strategy · Supplier collaboration · SAP · Power BI · Excel/VBA · "
            "Python · CNC programming"
        ),
        "experience": [
            ("2021 – 2026", "Wärtsilä", "Cost Engineer", [
                "Set the cost price for development, customer and non-standard projects — the "
                "reference figure engineering, sales, purchasing and business control all worked from.",
                "Translated technical choices into cost impact before decisions were made, as cost "
                "partner to engineering, sales, project management, purchasing and product management.",
                "Drove cross-functional cost-reduction and made budget-versus-actual transparent, "
                "so deviations surfaced early enough to act on.",
            ]),
            ("2019 – 2021", "Wilting", "Manufacturing Engineer", [
                "Built the workshop and tool management from the ground up and ran purchasing of "
                "indirect goods and tooling.",
                "Led improvement projects that raised safety, quality and delivery reliability "
                "while taking cost out of the process.",
            ]),
            ("2017 – 2019", "VDL ETG", "Factory Engineer", [
                "Defined make strategy and routing for complex high-tech products and produced the "
                "supporting manufacturing documentation.",
                "Owned running production as product owner — the link between design intent and "
                "what the factory actually makes.",
            ]),
            ("2011 – 2017", "Andritz Feed & Biofuel", "Production Supervisor", [
                "Led the hardening shop and three production departments on clear KPIs, with "
                "coaching and HR responsibility.",
                "Drove structural improvements in quality, throughput and cost across the "
                "departments.",
            ]),
            ("1987 – 2011", "DAF Trucks", "Production / Technical Engineer & Team Lead", [
                "Delivered process and product improvements, machine acceptance and the start-up "
                "of new production lines.",
                "Cut defects and improved line flow through continuous improvement and CNC "
                "programming; led and trained production teams.",
            ]),
        ],
        "education": ("KMBO Mechanical Engineering, Eindhoven  ·  "
                      "LTS Mechanical Engineering, Eindhoven"),
        "certs": "Lean Six Sigma — Green Belt & Yellow Belt",
        "languages": "Dutch (native) · English · German — spoken and written",
    },
    "nl": {
        "role": "Cost Engineer · Calculator · Werkvoorbereider",
        "headline": "Betrouwbare kostprijzen · snellere offertes · marges die kloppen",
        "headings": {"summary": "Profiel", "competencies": "Kerncompetenties",
                     "experience": "Werkervaring", "education": "Opleiding",
                     "certs": "Certificeringen", "languages": "Talen"},
        "summary": (
            "Cost engineer met een loopbaan van de werkvloer naar de calculatie — DAF Trucks, "
            "VDL ETG, Andritz, Wilting en Wärtsilä — zodat elke kostprijs die ik maak geworteld "
            "is in wat er werkelijk op de machine gebeurt. Ik bouw onderbouwde kostprijzen voor "
            "ontwikkel-, klant- en niet-standaard werk, houd offreren snel en transparant, en "
            "krijg engineering, inkoop, verkoop en business control op hetzelfde cijfer. Waar de "
            "tooling tekortschiet bouw ik die zelf — automatisering en dashboards in Python, en "
            "Excel/VBA. Lean Six Sigma Green Belt."
        ),
        "competencies": (
            "Kostencalculatie · Should-cost · RFQ & offerte · Nacalculatie · Make-or-buy · "
            "Stuklijst (BOM) · Marge- & kostenbeheersing · Value engineering · "
            "Lean Six Sigma (Green Belt) · Continu verbeteren · Werkvoorbereiding & routing · "
            "Maakstrategie · Leveranciers · SAP · Power BI · Excel/VBA · Python · CNC-programmeren"
        ),
        "experience": [
            ("2021 – 2026", "Wärtsilä", "Cost Engineer", [
                "De kostprijs bepaald voor ontwikkel-, klant- en niet-standaard projecten — het "
                "referentiecijfer waar engineering, verkoop, inkoop en business control allemaal "
                "mee werkten.",
                "Technische keuzes vertaald naar kosteneffect vóórdat beslissingen vielen, als "
                "kostenpartner voor engineering, verkoop, projectmanagement, inkoop en "
                "productmanagement.",
                "Cross-functionele kostenreductie aangejaagd en budget-versus-werkelijk "
                "transparant gemaakt, zodat afwijkingen vroeg genoeg zichtbaar werden om bij te "
                "sturen.",
            ]),
            ("2019 – 2021", "Wilting", "Manufacturing Engineer", [
                "Werkplaats en gereedschapsbeheer vanaf de basis opgezet en de inkoop van "
                "indirecte goederen en gereedschappen verzorgd.",
                "Verbeterprojecten geleid die veiligheid, kwaliteit en leverbetrouwbaarheid "
                "verhoogden en tegelijk kosten uit het proces haalden.",
            ]),
            ("2017 – 2019", "VDL ETG", "Factory Engineer", [
                "Maakstrategie en routing bepaald voor complexe high-tech producten en de "
                "bijbehorende productiedocumentatie opgesteld.",
                "Lopende productie als producteigenaar aangestuurd — de schakel tussen "
                "ontwerpintentie en wat de fabriek werkelijk maakt.",
            ]),
            ("2011 – 2017", "Andritz Feed & Biofuel", "Supervisor productie", [
                "De harderij en drie productieafdelingen aangestuurd op heldere KPI's, met "
                "coaching en HR-verantwoordelijkheid.",
                "Structurele verbeteringen in kwaliteit, doorlooptijd en kosten doorgevoerd over "
                "de afdelingen.",
            ]),
            ("1987 – 2011", "DAF Trucks", "Production / Technical Engineer & Teamleider", [
                "Proces- en productverbeteringen, machine-afnames en de opstart van nieuwe "
                "productielijnen gerealiseerd.",
                "Defecten teruggedrongen en de lijnflow verbeterd met continu verbeteren en "
                "CNC-programmeren; productieteams aangestuurd en opgeleid.",
            ]),
        ],
        "education": ("KMBO Metaaltechniek, Eindhoven  ·  "
                      "LTS Metaaltechniek, Eindhoven"),
        "certs": "Lean Six Sigma — Green Belt & Yellow Belt",
        "languages": "Nederlands (moedertaal) · Engels · Duits — in woord en geschrift",
    },
}

# Module-level aliases (English) used elsewhere in the app.
ROLE = CONTENT["en"]["role"]
KEYWORDS = CONTENT["en"]["competencies"]
PROFILE = CONTENT["en"]["summary"]
SKILLS = CONTENT["en"]["competencies"]


def build_pdf(path: Path, role_title: str | None = None, keywords: str | None = None, lang: str = "en") -> None:
    c = CONTENT[lang]
    role_title = role_title or c["role"]
    keywords = keywords or c["competencies"]
    h = c["headings"]

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9.5, leading=12.3,
                          textColor=INK, alignment=TA_LEFT)
    name = ParagraphStyle("name", parent=body, fontSize=21, leading=23, fontName="Helvetica-Bold", spaceAfter=2)
    role = ParagraphStyle("role", parent=body, fontSize=10.5, leading=14, textColor=ACCENT,
                          fontName="Helvetica-Bold", spaceAfter=2)
    headline = ParagraphStyle("headline", parent=body, fontSize=9.5, leading=12.5, textColor=INK,
                              fontName="Helvetica-Oblique", spaceAfter=3)
    contact = ParagraphStyle("contact", parent=body, fontSize=8.6, textColor=GREY, spaceAfter=2)
    h2 = ParagraphStyle("h2", parent=body, fontSize=10.5, leading=13, textColor=INK,
                        fontName="Helvetica-Bold", spaceBefore=7, spaceAfter=2)
    job = ParagraphStyle("job", parent=body, fontName="Helvetica-Bold", fontSize=10, spaceBefore=4)
    meta = ParagraphStyle("meta", parent=body, fontSize=8.6, textColor=GREY, spaceAfter=2)

    def heading(text):
        return [Paragraph(text.upper(), h2), HRFlowable(width="100%", thickness=0.6, color=RULE, spaceAfter=4)]

    def bullets(items):
        return ListFlowable([ListItem(Paragraph(t, body), leftIndent=10, value="•") for t in items],
                            bulletType="bullet", start="•", leftIndent=9, spaceBefore=1)

    flow = [Paragraph(NAME, name), Paragraph(role_title, role)]
    if c.get("headline"):
        flow.append(Paragraph(c["headline"], headline))
    flow += [Paragraph(CONTACT, contact),
             HRFlowable(width="100%", thickness=1.1, color=INK, spaceBefore=4, spaceAfter=2)]
    flow += heading(h["summary"]) + [Paragraph(c["summary"], body)]
    flow += heading(h["competencies"]) + [Paragraph(keywords, body)]
    flow += heading(h["experience"])
    for period, org, title, items in c["experience"]:
        flow += [Paragraph(title, job), Paragraph(f"{org}  ·  {period}", meta), bullets(items)]
    flow += heading(h["education"]) + [Paragraph(c["education"], body)]
    flow += heading(h["certs"]) + [Paragraph(c["certs"], body)]
    flow += heading(h["languages"]) + [Paragraph(c["languages"], body)]

    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
                      topMargin=13 * mm, bottomMargin=10 * mm, title=f"CV {NAME}", author=NAME).build(flow)


def build_docx(path: Path, role_title: str | None = None, keywords: str | None = None, lang: str = "en") -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    c = CONTENT[lang]
    role_title = role_title or c["role"]
    keywords = keywords or c["competencies"]
    h = c["headings"]

    doc = Document()
    base = doc.styles["Normal"].font
    base.name, base.size = "Calibri", Pt(10.5)

    def heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    n = doc.add_paragraph()
    rn = n.add_run(NAME)
    rn.bold = True
    rn.font.size = Pt(20)
    rr = doc.add_paragraph().add_run(role_title)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0x2A, 0x9D, 0x8F)
    if c.get("headline"):
        hp = doc.add_paragraph().add_run(c["headline"])
        hp.italic = True
        hp.font.size = Pt(10)
    doc.add_paragraph(CONTACT)

    heading(h["summary"]); doc.add_paragraph(c["summary"])
    heading(h["competencies"]); doc.add_paragraph(keywords)
    heading(h["experience"])
    for period, org, title, items in c["experience"]:
        jp = doc.add_paragraph()
        jr = jp.add_run(title)
        jr.bold = True
        mp = doc.add_paragraph().add_run(f"{org} · {period}")
        mp.italic = True
        mp.font.size = Pt(9)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")
    heading(h["education"]); doc.add_paragraph(c["education"])
    heading(h["certs"]); doc.add_paragraph(c["certs"])
    heading(h["languages"]); doc.add_paragraph(c["languages"])

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


if __name__ == "__main__":
    a = Path(__file__).parent / "assets"
    build_pdf(a / "cv.pdf", lang="en")
    build_docx(a / "cv.docx", lang="en")
    build_pdf(a / "cv_nl.pdf", lang="nl")
    build_docx(a / "cv_nl.docx", lang="nl")
    print(f"Wrote EN (cv.pdf/.docx) and NL (cv_nl.pdf/.docx) in {a}")
